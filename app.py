import base64
import datetime
import io
import json
import os
import tempfile
import fitz  # PyMuPDF
import google.generativeai as genai
import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageOps
import streamlit as st
import streamlit.components.v1 as components

APP_VERSION = "20260812-SECRETS-SAFE-FIX"

st.set_page_config(
    page_title=f"私車公用補助單自動化工具 ({APP_VERSION})", layout="centered"
)

st.caption(f"📌 程式版本：`{APP_VERSION}`")


# 注入 JavaScript：強效導航演算法，完美跳過 disabled 欄位並自動 Focus 下一個可輸入框
def inject_enter_focus_js():
    js_code = """
    <script>
    function setupEnterNavigation() {
        const doc = window.parent.document;
        
        function attachListeners() {
            const allInputs = Array.from(doc.querySelectorAll('input[type="text"], input[type="number"]'));
            
            allInputs.forEach((input) => {
                if (!input.dataset.enterBound) {
                    input.dataset.enterBound = "true";
                    input.addEventListener('keydown', function(e) {
                        if (e.key === 'Enter') {
                            e.preventDefault();
                            
                            const validInputs = Array.from(doc.querySelectorAll('input[type="text"]:not([disabled]), input[type="number"]:not([disabled])'));
                            const currentIndex = validInputs.indexOf(this);
                            
                            if (currentIndex !== -1 && currentIndex < validInputs.length - 1) {
                                const nextInput = validInputs[currentIndex + 1];
                                setTimeout(() => {
                                    nextInput.focus();
                                    if (typeof nextInput.select === 'function') {
                                        nextInput.select();
                                    }
                                }, 100);
                            }
                        }
                    });
                }
            });
        }

        attachListeners();
        setInterval(attachListeners, 1000);
    }
    
    setTimeout(setupEnterNavigation, 300);
    </script>
    """
    components.html(js_code, height=0, width=0)


# 注入右下角個人專屬署名 (Design by Max + 相容 .jpg/.jpeg/.png Q版頭像)
def inject_custom_footer():
    avatar_candidates = ["avatar.jpg", "avatar.jpeg", "avatar.png", "avatar.JPG"]
    img_base64 = ""
    mime_type = "image/png"

    for af in avatar_candidates:
        if os.path.exists(af):
            with open(af, "rb") as img_f:
                img_base64 = base64.b64encode(img_f.read()).decode("utf-8")
                if af.lower().endswith((".jpg", ".jpeg")):
                    mime_type = "image/jpeg"
                else:
                    mime_type = "image/png"
            break

    avatar_html = (
        f'<img src="data:{mime_type};base64,{img_base64}" style="width: 36px; height: 36px; border-radius: 50%; object-fit: cover; margin-right: 8px; border: 1.5px solid #ccc; background-color: #fff;">'
        if img_base64
        else ""
    )

    footer_css = f"""
    <style>
    .custom-footer-max {{
        position: fixed;
        bottom: 16px;
        right: 210px;
        display: flex;
        align-items: center;
        background-color: rgba(255, 255, 255, 0.95);
        padding: 4px 12px;
        border-radius: 20px;
        box-shadow: 0px 2px 8px rgba(0, 0, 0, 0.15);
        z-index: 999999;
        pointer-events: none;
    }}
    .custom-footer-text {{
        font-family: 'Comic Sans MS', cursive, sans-serif;
        font-weight: bold;
        font-style: italic;
        font-size: 0.95rem;
        color: #333333;
        white-space: nowrap;
    }}
    </style>
    <div class="custom-footer-max">
        {avatar_html}
        <span class="custom-footer-text">Design by Max</span>
    </div>
    """
    st.markdown(footer_css, unsafe_allow_html=True)


# 1. 網頁標題 (置中顯示、前後加車子圖示、單一行顯示)
st.markdown(
    "<h2 style='font-size: 1.6rem; font-weight: 700; white-space: nowrap; margin-top: -10px; text-align: center;'>🚗 私車公用補助單自動化填寫工具 🚗</h2>",
    unsafe_allow_html=True,
)

st.markdown(
    "上傳停車發票/收據照片或 **PDF 檔**，由 Gemini AI 自動辨識日期與金額，輕鬆生成報銷單！"
)

# 2. API Key 設定 (優先從 Streamlit Secrets 安全讀取)
default_key_from_secrets = st.secrets.get("GEMINI_API_KEY", "")

api_key = st.text_input(
    "請輸入 Gemini API Key：",
    type="password",
    value=default_key_from_secrets,
)

if not api_key:
    st.warning("⚠️ 請先在 Streamlit Secrets 設定 GEMINI_API_KEY 或於上方欄位輸入 API Key。")
    st.stop()

# 3. 基本資料填寫 (預設全為空字串，以 placeholder 提供填寫提示)
col1, col2 = st.columns(2)
with col1:
    user_name = st.text_input("姓名", value="", placeholder="例如：王大明")
with col2:
    user_dept = st.text_input("部門", value="", placeholder="例如：研發部")

# 4. 上傳檔案
uploaded_files = st.file_uploader(
    "1. 上傳停車發票/收據（照片或 PDF 檔，可多選）",
    type=["jpg", "jpeg", "png", "pdf"],
    accept_multiple_files=True,
)


def format_date_to_excel(d_str):
    """將 YYYYMMDD 純數字格式化為標準日期 YYYY/MM/DD"""
    d_str = str(d_str).strip()
    if len(d_str) == 8 and d_str.isdigit():
        return f"{d_str[:4]}/{d_str[4:6]}/{d_str[6:]}"
    return d_str


def set_cell_value(ws, cell_ref, value):
    """安全寫入函式：自動相容一般儲存格與合併儲存格 (MergedCell)"""
    if isinstance(cell_ref, str):
        cell = ws[cell_ref]
    else:
        cell = ws.cell(row=cell_ref[0], column=cell_ref[1])

    if type(cell).__name__ == "MergedCell":
        for rng in ws.merged_cells.ranges:
            if cell.coordinate in rng:
                ws.cell(row=rng.min_row, column=rng.min_col).value = value
                return
    cell.value = value


def crop_and_rotate_receipt_bytes(raw_bytes, box_2d, rotate_deg):
    """將照片轉正 (EXIF 修正)、去背裁切，並旋轉為直立長條方向"""
    try:
        image = Image.open(io.BytesIO(raw_bytes))
        image = ImageOps.exif_transpose(image)

        if box_2d and isinstance(box_2d, list) and len(box_2d) == 4:
            w, h = image.size
            ymin, xmin, ymax, xmax = box_2d

            left = max(0, int(xmin * w / 1000) - 10)
            top = max(0, int(ymin * h / 1000) - 10)
            right = min(w, int(xmax * w / 1000) + 10)
            bottom = min(h, int(ymax * h / 1000) + 10)

            if right > left and bottom > top:
                image = image.crop((left, top, right, bottom))

        if rotate_deg in [90, 180, 270]:
            if rotate_deg == 90:
                image = image.transpose(Image.ROTATE_270)
            elif rotate_deg == 180:
                image = image.transpose(Image.ROTATE_180)
            elif rotate_deg == 270:
                image = image.transpose(Image.ROTATE_90)

        out_io = io.BytesIO()
        image.save(out_io, format="PNG")
        return out_io.getvalue()
    except Exception:
        return raw_bytes


def process_images_with_gemini(files, key):
    genai.configure(api_key=key.strip())

    candidate_models = [
        "gemini-1.5-flash",
        "gemini-1.5-pro",
        "gemini-2.0-flash",
        "gemini-2.5-flash",
    ]

    try:
        available = [
            m.name.replace("models/", "")
            for m in genai.list_models()
            if "generateContent" in m.supported_generation_methods
        ]
        for a in available:
            if a not in candidate_models:
                candidate_models.insert(0, a)
    except Exception:
        pass

    prompt = """
    你是一個財務報銷助手。請讀取這份發票/收據照片或文件，並提取以下資訊：
    1. date: 發票日期，格式請統一轉換為 YYYYMMDD（例如 20260603）。若無法取得年份，預設為當前年份。
    2. amount: 停車費金額 (數字)。
    3. box_2d: 圖片中「收據/發票紙張本體」的範圍座標 [ymin, xmin, ymax, xmax]，數值請以 0 到 1000 之間的整數表示。請貼緊收據邊緣去背。若為 PDF 則輸出 [0, 0, 1000, 1000]。
    4. rotate: 圖片中收據文字的方向。為了讓收據變成「文字由左至右、由上至下正向讀取」的直立長條狀，請判斷需要【順時針旋轉多少度】：
       - 若文字已經正面朝上：輸出 0
       - 若文字頭朝左：輸出 90
       - 若文字倒立 (頭朝下)：輸出 180
       - 若文字頭朝右：輸出 270
    
    請直接輸出純 JSON 格式，例如：
    {"date": "20260603", "amount": 150, "box_2d": [150, 250, 850, 750], "rotate": 0}
    注意：絕對不要加上 ```json 或任何 markdown 標記。
    """

    results = []
    for uploaded_file in files:
        bytes_data = uploaded_file.read()
        file_ext = uploaded_file.name.split(".")[-1].lower()

        if file_ext == "pdf":
            mime_type = "application/pdf"
        elif file_ext == "png":
            mime_type = "image/png"
        else:
            mime_type = "image/jpeg"

        content_part = {"mime_type": mime_type, "data": bytes_data}

        success = False
        last_error = ""

        for model_name in candidate_models:
            try:
                model = genai.GenerativeModel(model_name)
                response = model.generate_content([prompt, content_part])

                raw_text = response.text.strip()
                clean_text = (
                    raw_text.replace("```json", "")
                    .replace("```", "")
                    .strip()
                )

                res_json = json.loads(clean_text)

                box_2d = res_json.get("box_2d", [0, 0, 1000, 1000])
                rotate_deg = res_json.get("rotate", 0)

                if file_ext in ["jpg", "jpeg", "png"]:
                    processed_bytes = crop_and_rotate_receipt_bytes(
                        bytes_data, box_2d, rotate_deg
                    )
                else:
                    processed_bytes = bytes_data

                res_json["raw_bytes"] = processed_bytes
                res_json["file_ext"] = file_ext
                res_json["filename"] = uploaded_file.name
                results.append(res_json)
                success = True
                break
            except Exception as e:
                last_error = str(e)
                continue

        if not success:
            st.error(
                f"❌ 解析檔案 『{uploaded_file.name}』 失敗：{last_error}"
            )
            st.stop()

    results.sort(key=lambda x: x["date"])
    return results


if uploaded_files and user_name and user_dept:
    if st.button("🤖 AI 辨識單據內容"):
        with st.spinner("Gemini 分析照片/PDF 中..."):
            st.session_state["parsed_receipts"] = process_images_with_gemini(
                uploaded_files, api_key
            )
        st.success(
            f"成功辨識 {len(st.session_state['parsed_receipts'])} 筆單據資料！"
        )

# 5. 補充欄位與匯出 Excel / Word
if "parsed_receipts" in st.session_state:
    receipts = st.session_state["parsed_receipts"]

    st.subheader("📝 補充填寫報銷明細")

    chk_col1, chk_col2, chk_col3, chk_col4 = st.columns(4)
    same_loc = chk_col1.checkbox("所有【地點】相同")
    same_km = chk_col2.checkbox("所有【公里數】相同")
    same_toll = chk_col3.checkbox("所有【過路費】相同")
    same_reason = chk_col4.checkbox("所有【事由】相同")

    st.markdown("---")
    details = []

    first_loc, first_km, first_toll, first_reason = "", 0, 0, ""

    for idx, r in enumerate(receipts):
        formatted_date = format_date_to_excel(r["date"])
        st.markdown(
            f"**單據 {idx+1} (日期: {formatted_date} / 停車費: {r['amount']}元)**"
        )
        c1, c2, c3, c4 = st.columns(4)

        # 1. 地點欄位
        if idx == 0:
            loc_val = c1.text_input(
                f"地點 #{idx+1}",
                value="",
                placeholder="例如：客戶端",
                key=f"loc_{idx}",
            )
            first_loc = loc_val
        else:
            if same_loc:
                loc_val = first_loc
                c1.text_input(
                    f"地點 #{idx+1}",
                    value=first_loc,
                    disabled=True,
                    key=f"dis_loc_{idx}",
                )
            else:
                loc_val = c1.text_input(
                    f"地點 #{idx+1}",
                    value="",
                    placeholder="例如：客戶端",
                    key=f"loc_{idx}",
                )

        # 2. 公里數欄位
        if idx == 0:
            km_val = int(
                c2.number_input(
                    f"公里數 #{idx+1}", value=0, step=1, key=f"km_{idx}"
                )
            )
            first_km = km_val
        else:
            if same_km:
                km_val = int(first_km)
                c2.number_input(
                    f"公里數 #{idx+1}",
                    value=int(first_km),
                    disabled=True,
                    key=f"dis_km_{idx}",
                )
            else:
                km_val = int(
                    c2.number_input(
                        f"公里數 #{idx+1}", value=0, step=1, key=f"km_{idx}"
                    )
                )

        # 3. 回數票/過路費欄位
        if idx == 0:
            toll_val = int(
                c3.number_input(
                    f"回數票 #{idx+1}", value=0, step=1, key=f"toll_{idx}"
                )
            )
            first_toll = toll_val
        else:
            if same_toll:
                toll_val = int(first_toll)
                c3.number_input(
                    f"回數票 #{idx+1}",
                    value=int(first_toll),
                    disabled=True,
                    key=f"dis_toll_{idx}",
                )
            else:
                toll_val = int(
                    c3.number_input(
                        f"回數票 #{idx+1}", value=0, step=1, key=f"toll_{idx}"
                    )
                )

        # 4. 事由欄位
        if idx == 0:
            reason_val = c4.text_input(
                f"事由 #{idx+1}",
                value="",
                placeholder="例如：拜訪客戶",
                key=f"reason_{idx}",
            )
            first_reason = reason_val
        else:
            if same_reason:
                reason_val = first_reason
                c4.text_input(
                    f"事由 #{idx+1}",
                    value=first_reason,
                    disabled=True,
                    key=f"dis_reason_{idx}",
                )
            else:
                reason_val = c4.text_input(
                    f"事由 #{idx+1}",
                    value="",
                    placeholder="例如：拜訪客戶",
                    key=f"reason_{idx}",
                )

        details.append(
            {
                "date": formatted_date,
                "location": loc_val,
                "km": km_val,
                "parking": int(round(float(r["amount"]))),
                "toll": toll_val,
                "reason": reason_val,
                "raw_bytes": r["raw_bytes"],
                "file_ext": r["file_ext"],
            }
        )

    st.markdown(" ")
    btn_col1, btn_col2 = st.columns(2)

    # 產出 Excel
    with btn_col1:
        if st.button("🚀 產出 Excel 報銷檔案"):
            template_xlsx = "私車公用補助申請單.xlsx"

            if not os.path.exists(template_xlsx):
                st.error(
                    "系統找不到範本『私車公用補助申請單.xlsx』！請確認檔名是否包含 .xlsx"
                )
            else:
                wb = openpyxl.load_workbook(template_xlsx)

                # Sheet 1: 私車公用單
                ws1 = wb.worksheets[0]

                set_cell_value(ws1, "B3", user_name)
                set_cell_value(ws1, "E3", user_dept)

                for i, item in enumerate(details):
                    row_num = 5 + i
                    set_cell_value(ws1, (row_num, 1), item["date"])
                    set_cell_value(ws1, (row_num, 2), item["location"])
                    set_cell_value(ws1, (row_num, 4), item["km"])
                    set_cell_value(ws1, (row_num, 5), item["parking"])
                    set_cell_value(ws1, (row_num, 6), item["toll"])
                    set_cell_value(ws1, (row_num, 8), item["reason"])

                # Sheet 2: 支出憑單
                ws2 = wb.worksheets[1]
                today_str = datetime.datetime.now().strftime("%Y年%m月%d日")

                set_cell_value(ws2, "G5", today_str)
                set_cell_value(ws2, "C7", f"專案編號：{user_name}")

                first_date = details[0]["date"]
                last_date = details[-1]["date"]
                set_cell_value(ws2, "A9", f"{first_date}~{last_date}交通費用")

                tot_km = sum(item["km"] for item in details)
                tot_parking = sum(item["parking"] for item in details)
                tot_toll = sum(item["toll"] for item in details)
                grand_total = (tot_km * 6) + tot_parking + tot_toll

                set_cell_value(ws2, "G9", grand_total)
                set_cell_value(ws2, "G17", grand_total)

                output_date = datetime.datetime.now().strftime("%Y%m%d")
                out_filename = f"私車公用補助申請單-{user_name}-{output_date}.xlsx"

                with tempfile.NamedTemporaryFile(
                    delete=False, suffix=".xlsx"
                ) as tmp:
                    wb.save(tmp.name)
                    tmp_path = tmp.name

                with open(tmp_path, "rb") as file:
                    st.download_button(
                        label="📥 下載報銷單 (Excel)",
                        data=file,
                        file_name=out_filename,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )

    # 產出 Word 報支單據憑證檔 (強行同頁 + 動態等比縮放防跨頁)
    with btn_col2:
        if st.button("📄 產出 Word 報支單據檔"):
            doc = Document()

            for idx, item in enumerate(details):
                if idx > 0:
                    doc.add_page_break()

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0

                # 1. 寫入標題文字
                run_title = p.add_run(f"{item['date']} 停車費\n\n")
                run_title.font.size = Pt(14)
                run_title.font.bold = True
                run_title.font.name = "標楷體"
                run_title._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

                # 2. 寫入圖片（動態計算比例防跨頁）
                raw_bytes = item["raw_bytes"]
                ext = item["file_ext"]

                img_stream = None
                if ext == "pdf":
                    pdf_doc = fitz.open(stream=raw_bytes, filetype="pdf")
                    if len(pdf_doc) > 0:
                        page = pdf_doc[0]
                        pix = page.get_pixmap(dpi=200)
                        img_stream = io.BytesIO(pix.tobytes("png"))
                else:
                    img_stream = io.BytesIO(raw_bytes)

                if img_stream:
                    try:
                        pil_img = Image.open(img_stream)
                        w, h = pil_img.size
                        aspect_ratio = h / w if w > 0 else 1.0

                        target_width = 2.2
                        target_height = target_width * aspect_ratio

                        if target_height > 4.8:
                            target_height = 4.8
                            target_width = target_height / aspect_ratio

                        img_stream.seek(0)
                        run_img = p.add_run()
                        run_img.add_picture(
                            img_stream, width=Inches(target_width)
                        )
                    except Exception as e:
                        p.add_run(f"[圖片載入失敗: {e}]")

            output_date = datetime.datetime.now().strftime("%Y%m%d")
            word_filename = f"報支單據-{user_name}-{output_date}.docx"

            with tempfile.NamedTemporaryFile(
                delete=False, suffix=".docx"
            ) as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name

            with open(tmp_path, "rb") as file:
                st.download_button(
                    label="📥 下載報支單據 (Word)",
                    data=file,
                    file_name=word_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

inject_enter_focus_js()
inject_custom_footer()
